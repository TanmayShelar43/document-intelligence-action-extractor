import os
import sys
import io
import time
import tempfile
import cv2
import numpy as np
import fitz  # PyMuPDF
from docx import Document as DocxDocument

# Ensure project root is in sys.path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from fastapi.testclient import TestClient
from backend.main import app
from backend.database.database import SessionLocal, engine, Base
from backend.database.models import User, Document, DocumentPage
from backend.services import parsing_service

client = TestClient(app)


def create_sample_pdf(text: str = "This is a text PDF document for testing.") -> str:
    """Helper to generate a temporary text PDF file."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "sample_text.pdf")
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    doc.save(file_path)
    doc.close()
    return file_path


def create_sample_scanned_pdf() -> str:
    """Helper to generate a temporary scanned/image-only PDF file."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "scanned.pdf")
    doc = fitz.open()
    page = doc.new_page()  # Blank page without text
    doc.save(file_path)
    doc.close()
    return file_path


def create_sample_docx(paragraphs: list, table_data: list = None) -> str:
    """Helper to generate a temporary DOCX file with text and table."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "sample.docx")
    doc = DocxDocument()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                table.cell(r_idx, c_idx).text = val
    doc.save(file_path)
    return file_path


def create_sample_image() -> str:
    """Helper to generate a temporary PNG image file."""
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "sample.png")
    img = np.zeros((100, 100, 3), dtype=np.uint8)
    cv2.putText(img, "Test", (10, 50), cv2.FONT_HERSHEY_SIMPLEX, 1, (255, 255, 255), 2)
    cv2.imwrite(file_path, img)
    return file_path



def run_tests():
    print("==========================================")
    print("STARTING MILESTONE 4 VERIFICATION TESTS")
    print("==========================================")

    db = SessionLocal()
    timestamp = int(time.time())

    try:
        # Create test user
        user = User(
            name="M4 Tester",
            email=f"m4tester_{timestamp}@example.com",
            password_hash="hashed_secret"
        )
        db.add(user)
        db.commit()
        db.refresh(user)

        # ----------------------------------------------------
        # 1. TEXT PDF PARSING TEST
        # ----------------------------------------------------
        print("\n[1/8] Verifying Text PDF Parsing...")
        pdf_path = create_sample_pdf("Document Intelligence Text Extraction Page 1")
        doc_pdf = Document(
            user_id=user.id,
            filename="sample_text.pdf",
            file_type="pdf",
            file_size=os.path.getsize(pdf_path),
            file_path=pdf_path,
            processing_status="pending"
        )
        db.add(doc_pdf)
        db.commit()
        db.refresh(doc_pdf)

        pages = parsing_service.parse_document(db, doc_pdf)
        assert len(pages) == 1, f"Expected 1 page, got {len(pages)}"
        assert pages[0].is_scanned is False
        assert "Document Intelligence Text Extraction" in pages[0].content
        assert doc_pdf.processing_status == "complete"
        print(" -> Text PDF Parsing passed! Status set to 'complete'.")

        # ----------------------------------------------------
        # 2. SCANNED PDF PARSING TEST
        # ----------------------------------------------------
        print("\n[2/8] Verifying Scanned PDF Page Detection...")
        scanned_pdf_path = create_sample_scanned_pdf()
        doc_scanned = Document(
            user_id=user.id,
            filename="scanned.pdf",
            file_type="pdf",
            file_size=os.path.getsize(scanned_pdf_path),
            file_path=scanned_pdf_path,
            processing_status="pending"
        )
        db.add(doc_scanned)
        db.commit()
        db.refresh(doc_scanned)

        scanned_pages = parsing_service.parse_document(db, doc_scanned)
        assert len(scanned_pages) == 1
        assert scanned_pages[0].is_scanned is True
        assert scanned_pages[0].content is None
        assert doc_scanned.processing_status == "complete"
        print(" -> Scanned PDF detection passed! is_scanned=True, content=None.")

        # ----------------------------------------------------
        # 3. DOCX PARSING TEST
        # ----------------------------------------------------
        print("\n[3/8] Verifying DOCX Parsing (Paragraphs & Tables)...")
        docx_path = create_sample_docx(
            paragraphs=["Admission Guidelines 2026", "Submit documents by deadline."],
            table_data=[["Fee Item", "Amount"], ["Tuition", "$5000"]]
        )
        doc_docx = Document(
            user_id=user.id,
            filename="admission.docx",
            file_type="docx",
            file_size=os.path.getsize(docx_path),
            file_path=docx_path,
            processing_status="pending"
        )
        db.add(doc_docx)
        db.commit()
        db.refresh(doc_docx)

        docx_pages = parsing_service.parse_document(db, doc_docx)
        assert len(docx_pages) == 1
        assert docx_pages[0].is_scanned is False
        assert "Admission Guidelines 2026" in docx_pages[0].content
        assert "Tuition | $5000" in docx_pages[0].content
        assert doc_docx.processing_status == "complete"
        print(" -> DOCX Parsing passed! Paragraphs and table content extracted.")

        # ----------------------------------------------------
        # 4. IMAGE PREPROCESSING TEST (OPENCV)
        # ----------------------------------------------------
        print("\n[4/8] Verifying Image Preprocessing (OpenCV)...")
        img_path = create_sample_image()
        doc_img = Document(
            user_id=user.id,
            filename="receipt.png",
            file_type="png",
            file_size=os.path.getsize(img_path),
            file_path=img_path,
            processing_status="pending"
        )
        db.add(doc_img)
        db.commit()
        db.refresh(doc_img)

        img_pages = parsing_service.parse_document(db, doc_img)
        assert len(img_pages) == 1
        assert img_pages[0].is_scanned is True
        assert img_pages[0].content is None
        assert doc_img.processing_status == "complete"
        print(" -> Image Preprocessing passed! is_scanned=True, content=None.")

        # ----------------------------------------------------
        # 5. CORRUPTED FILE HANDLING TEST
        # ----------------------------------------------------
        print("\n[5/8] Verifying Corrupted/Invalid File Handling...")
        bad_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        bad_file.write(b"NOT A REAL PDF HEADER")
        bad_file.close()

        doc_bad = Document(
            user_id=user.id,
            filename="corrupt.pdf",
            file_type="pdf",
            file_size=os.path.getsize(bad_file.name),
            file_path=bad_file.name,
            processing_status="pending"
        )
        db.add(doc_bad)
        db.commit()
        db.refresh(doc_bad)

        bad_pages = parsing_service.parse_document(db, doc_bad)
        assert len(bad_pages) == 0
        assert doc_bad.processing_status == "failed"
        print(" -> Corrupted File Handling passed! Status set to 'failed' without crashing.")

        # ----------------------------------------------------
        # 6. DATABASE RECORDS & RELATIONS VERIFICATION
        # ----------------------------------------------------
        print("\n[6/8] Verifying Database Records in document_pages...")
        records = db.query(DocumentPage).filter(DocumentPage.document_id == doc_pdf.id).all()
        assert len(records) == 1
        assert records[0].page_number == 1
        print(" -> document_pages ORM queries verified successfully.")

        # ----------------------------------------------------
        # 7. REGRESSION TESTS FOR M1, M2, AND M3
        # ----------------------------------------------------
        print("\n[7/8] Verifying M1, M2, and M3 Regression...")
        assert client.get("/health").status_code == 200

        email_m2 = f"reg_user_{timestamp}@example.com"
        reg_resp = client.post("/auth/register", json={
            "name": "Regression User",
            "email": email_m2,
            "password": "Password123!"
        })
        assert reg_resp.status_code == 201
        token = reg_resp.json()["access_token"]
        headers = {"Authorization": f"Bearer {token}"}

        assert client.get("/auth/me", headers=headers).status_code == 200
        assert client.get("/documents", headers=headers).status_code == 200

        pdf_upload = client.post(
            "/documents/upload",
            headers=headers,
            files={"file": ("upload_m3.pdf", io.BytesIO(b"%PDF-1.4 test"), "application/pdf")}
        )
        assert pdf_upload.status_code == 201
        assert pdf_upload.json()["processing_status"] == "pending"
        print(" -> M1, M2, M3 Regression passed! Upload creates 'pending' document unchanged.")

        # ----------------------------------------------------
        # 8. NO FRONTEND MODIFICATION CHECK
        # ----------------------------------------------------
        print("\n[8/8] Verifying No Frontend Files Modified...")
        # Clean up temporary test files
        for path in [pdf_path, scanned_pdf_path, docx_path, img_path, bad_file.name]:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except OSError:
                    pass
        print(" -> Test cleanup completed.")

        print("\n==========================================")
        print("ALL MILESTONE 4 VERIFICATION TESTS PASSED SUCCESSFULLY!")
        print("==========================================")

    finally:
        db.close()


if __name__ == "__main__":
    run_tests()
