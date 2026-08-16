from typing import List, Dict, Any
from docx import Document as DocxDocument


def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Parses DOCX/DOC files using python-docx.
    Extracts paragraphs, headings, and tables into structured text content.
    Returns page data dictionary list.
    """
    doc = DocxDocument(file_path)
    text_blocks = []

    # Extract paragraphs and headings
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_blocks.append(text)

    # Extract tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                table_rows.append(row_text)
        if table_rows:
            text_blocks.append("\n".join(table_rows))

    combined_content = "\n\n".join(text_blocks).strip()

    return [{
        "page_number": 1,
        "content": combined_content if combined_content else None,
        "is_scanned": False
    }]
